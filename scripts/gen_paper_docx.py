"""Generate a DOCX export of the paper from markdown sections."""
import os
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "--quiet"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SECTIONS_DIR = os.path.join(ROOT, 'public', 'data', 'sections')
OUT_PATH = os.path.join(ROOT, 'public', 'paper.docx')

md_files = [
    'abstract.md', 'introduction.md', 'methodology.md',
    'experiments.md', 'discussion.md', 'conclusion.md', 'references.md'
]

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

def add_markdown(doc, text):
    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith('# ') and not stripped.startswith('## '):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith('## ') and not stripped.startswith('### '):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith('### ') and not stripped.startswith('#### '):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif stripped.startswith('$$') and stripped.endswith('$$'):
            formula = stripped.strip('$$').strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(formula)
            run.italic = True
            run.font.size = Pt(10)
        elif stripped.startswith('|'):
            table_lines = [stripped]
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith('|'):
                table_lines.append(lines[j].strip())
                j += 1
            i = j - 1
            rows = []
            for tl in table_lines:
                if '---' in tl:
                    continue
                cells = [c.strip() for c in tl.strip('|').split('|')]
                rows.append(cells)
            if rows:
                max_cols = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=max_cols)
                tbl.style = 'Table Grid'
                for ri, row in enumerate(rows):
                    for ci, cell_text in enumerate(row):
                        if ci < max_cols:
                            tbl.rows[ri].cells[ci].text = cell_text
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            text_with_math = re.sub(r'\$([^$]+)\$', r'\(\1\)', stripped)
            run = p.add_run(text_with_math)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        i += 1

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run('RF Spectrogram Anomaly Detection with Quantum Kitchen Sinks')
run.bold = True
run.font.size = Pt(18)
run.font.name = 'Times New Roman'

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Abdallah Aaraba, Alexis Vieloszynski, Remon Polus, Ola Ahmad, Soumaya Cherkaoui\n')
meta.add_run('Polytechnique Montr\u00e9al \u00b7 Thales cortAIx Lab')
meta.runs[0].font.size = Pt(10)
meta.runs[1].font.size = Pt(10)

doc.add_paragraph()  # spacer

for fname in md_files:
    fpath = os.path.join(SECTIONS_DIR, fname)
    if os.path.exists(fpath):
        text = open(fpath, encoding='utf-8').read()
        add_markdown(doc, text)
        if fname != md_files[-1]:
            doc.add_page_break()

os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)
print(f'OK -> {OUT_PATH}')
